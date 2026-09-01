from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "IngredientIQ_16_Week_Project_Plan.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(paragraph, before=0, after=6, line=1.10, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    return run


def add_text_paragraph(doc, text, style=None, before=0, after=6, color=BLACK, size=11, bold_prefix=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after, line=1.10, keep_with_next=False)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
        set_paragraph_format(p, before=16, after=8, line=1.10, keep_with_next=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
        set_paragraph_format(p, before=12, after=6, line=1.10, keep_with_next=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
        set_paragraph_format(p, before=8, after=4, line=1.10, keep_with_next=True)
    return p


def add_callout(doc, heading, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=180)
    remove_table_borders(table)
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.10)
    r = p.add_run(heading + " ")
    set_run_font(r, size=10.5, color=INK, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, before=0, after=4, line=1.0)


def add_role_line(doc, label, task):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_format(p, before=0, after=3, line=1.10)
    r = p.add_run(label + ": ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(task)
    set_run_font(r, size=10.5, color=BLACK)


def add_done_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_format(p, before=1, after=7, line=1.10)
    r = p.add_run("Definition of done: ")
    set_run_font(r, size=10.2, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK, italic=True)


def add_week(doc, number, title, a, b, c, done):
    add_heading(doc, f"Week {number}: {title}", level=2)
    add_role_line(doc, "Person A - Frontend", a)
    add_role_line(doc, "Person B - Backend / Database", b)
    add_role_line(doc, "Person C - Data / Research", c)
    add_done_line(doc, done)


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def configure_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(header_p, before=0, after=0, line=1.0)
    run = header_p.add_run("INGREDIENTIQ | 16-WEEK PROJECT PLAN")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(footer_p, before=0, after=0, line=1.0)
    run = footer_p.add_run("Working plan | Page ")
    set_run_font(run, size=8.5, color=MUTED)
    page_run = add_page_field(footer_p)
    set_run_font(page_run, size=8.5, color=MUTED)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=28, after=5, line=1.0, keep_with_next=True)
    r = p.add_run("INGREDIENTIQ")
    set_run_font(r, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=7, line=1.0, keep_with_next=True)
    r = p.add_run("A Practical 16-Week Working Plan")
    set_run_font(r, size=16, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=22, line=1.0)
    r = p.add_run("Three-person web-game project | Four-month execution schedule")
    set_run_font(r, size=11, color=MUTED, italic=True)

    add_callout(
        doc,
        "Project framing.",
        "Build a text-first web game in which players guess a food item's NOVA processing level (1-4). The game uses a locally curated food-data snapshot, not machine learning or live API calls during gameplay.",
    )


def add_roles_section(doc):
    add_heading(doc, "Team roles", level=1)
    add_text_paragraph(
        doc,
        "Each person owns a primary area, but all three join the weekly demo, card review, and final testing.",
        before=0,
        after=8,
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1440, 3300, 4620], indent=120)
    mark_header_row(table.rows[0])
    header = table.rows[0].cells
    for cell, text in zip(header, ["Person", "Primary ownership", "Weekly responsibility"]):
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=0, after=0, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)
    roles = [
        ("A", "Frontend and game UX", "Game screens, interactions, responsive design, and usability fixes."),
        ("B", "Backend, database, deployment", "Food/session/response data, game endpoints, exports, and hosting."),
        ("C", "Data, research, testing", "Card curation, methodology, user testing, pilot coordination, and analysis."),
    ]
    for person, ownership, responsibility in roles:
        cells = table.add_row().cells
        for cell, value in zip(cells, [f"Person {person}", ownership, responsibility]):
            p = cell.paragraphs[0]
            set_paragraph_format(p, before=0, after=0, line=1.08)
            r = p.add_run(value)
            set_run_font(r, size=10, color=BLACK, bold=(cell == cells[0]))
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, before=0, after=3, line=1.0)


def add_scope_section(doc):
    add_heading(doc, "Scope rules that keep the project realistic", level=1)
    add_callout(
        doc,
        "Use the right term.",
        "The available dataset field is nova_group, which represents NOVA processing level. Do not present it as a measured preservative level unless a separate, documented preservative rubric is introduced.",
    )
    items = [
        "Use only manually approved food cards in the game. Start with 100-150 cards, balanced across Levels 1-4.",
        "Use anonymous session IDs first. Add accounts or login only if the professor specifically requires them.",
        "Do not add machine learning, a live Open Food Facts API dependency, product-image scraping, a leaderboard, or a full admin panel before the MVP works.",
        "If personal information is collected or a formal study is required, confirm consent and ethics requirements with the professor before recruiting participants.",
    ]
    for item in items:
        add_text_paragraph(doc, item, before=0, after=4, size=10.5, bold_prefix=None)


def add_phase_overview(doc):
    add_heading(doc, "Four-month overview", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1200, 3900, 4260], indent=120)
    mark_header_row(table.rows[0])
    headers = ["Weeks", "Focus", "Visible outcome"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=0, after=0, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)
    rows = [
        ("1-4", "Foundation and core game", "A real, local 10-question game using approved food data."),
        ("5-8", "MVP completion and internal testing", "Mobile-ready MVP and a shareable staging link."),
        ("9-12", "Pilot preparation and data collection", "Stable pilot with anonymous, usable response data."),
        ("13-16", "Analysis, report, and final demo", "Results, documentation, polished deployment, and presentation."),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row_values)):
            p = cell.paragraphs[0]
            set_paragraph_format(p, before=0, after=0, line=1.08)
            r = p.add_run(value)
            set_run_font(r, size=10, color=BLACK, bold=(index == 0))
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, before=0, after=2, line=1.0)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc.sections[0])
    set_document_styles(doc)

    add_title_block(doc)
    add_roles_section(doc)
    add_scope_section(doc)
    add_phase_overview(doc)

    doc.add_page_break()
    add_heading(doc, "Month 1: Foundation and core game", level=1)
    add_week(
        doc,
        1,
        "Scope and data setup",
        "Map one complete round: food card -> answer -> reveal -> score. Agree on the minimum screens.",
        "Choose the stack, create the project repository, and define the foods, sessions, and responses data fields.",
        "Confirm the NOVA terminology, inspect the source dataset, and document the card-approval rules.",
        "The team agrees on the game flow, data fields, and minimum feature list.",
    )
    add_week(
        doc,
        2,
        "Prepare approved game cards and a static prototype",
        "Build one static 10-question page with four answer buttons and a score counter using temporary cards.",
        "Create the database schema and a repeatable import path for approved food cards.",
        "Lead review of the 200 candidate cards. Split the work across the team and approve or reject cards with notes.",
        "At least 100 approved cards exist, with at least 25 per NOVA level, and the static prototype can be clicked through.",
    )
    add_week(
        doc,
        3,
        "Connect real data to the game",
        "Replace temporary cards with real data and show one random card at a time.",
        "Seed the approved foods table and build the endpoint that returns a question without exposing the answer first.",
        "Check card balance, ingredient text, names, duplicates, and unclear examples. Finalize the initial data dictionary.",
        "A player can complete a local round of 10 real food questions.",
    )
    add_week(
        doc,
        4,
        "Finish the first complete game loop",
        "Add answer feedback, reveal content, a score summary, and restart behaviour.",
        "Save guessed level, actual level, response time, and anonymous session ID for every answer.",
        "Write short and factual explanation text; do a full content check of every card currently seeded.",
        "The team can demonstrate an end-to-end 10-question game and confirm that response records are saved correctly.",
    )

    doc.add_page_break()
    add_heading(doc, "Month 2: MVP completion and internal testing", level=1)
    add_week(
        doc,
        5,
        "Make the MVP reliable on mobile",
        "Improve responsive layout, loading states, and clear instructions for first-time players.",
        "Prevent repeated food cards within a round and validate all answer logging rules.",
        "Replace weak or unfamiliar cards and maintain a balanced approved pool of at least 100 cards.",
        "The game works cleanly on a phone and desktop browser with real cards.",
    )
    add_week(
        doc,
        6,
        "Add minimal data collection and results support",
        "Improve the start, results, and replay screens. Keep the design simple and readable.",
        "Add anonymous sessions and a basic response export for later analysis. Do not build login unless required.",
        "Define the metrics to report: accuracy by level, common wrong guesses, score, and response time.",
        "A test round creates an exportable record with all research-relevant fields.",
    )
    add_week(
        doc,
        7,
        "Deploy a staging version and test usability",
        "Polish accessibility, button labels, and mobile spacing. Fix only issues found in testing.",
        "Deploy a staging link and confirm the data path works outside the local computer.",
        "Run a 5-10 person usability test and collect specific feedback on unclear cards or instructions.",
        "A shareable staging link works and has a short prioritized feedback list.",
    )
    add_week(
        doc,
        8,
        "MVP checkpoint",
        "Fix important usability issues; avoid adding new game modes or visual features.",
        "Fix data, deployment, and logging bugs discovered during usability testing.",
        "Update the card list and wording based on feedback. Prepare a brief checkpoint summary for the professor.",
        "The MVP is stable enough for a professor demo and pilot preparation.",
    )

    doc.add_page_break()
    add_heading(doc, "Month 3: Pilot preparation and data collection", level=1)
    add_week(
        doc,
        9,
        "Prepare the pilot without expanding scope",
        "Make only small clarity improvements and create a short in-game instruction screen.",
        "Add a simple safe import or update workflow only if changing approved cards is otherwise difficult.",
        "Prepare participant instructions, a recruitment message, and consent text if required by the professor.",
        "The team has a stable pilot build and a clear participant workflow.",
    )
    add_week(
        doc,
        10,
        "Launch the pilot",
        "Support participants and fix only game-breaking problems.",
        "Monitor errors and validate that new response records are being stored and exported correctly.",
        "Begin recruitment and track participation. Record any card complaints or obvious classification issues.",
        "The first real participant responses are recorded successfully.",
    )
    add_week(
        doc,
        11,
        "Continue data collection",
        "Keep the user experience stable; do not change game logic mid-pilot unless it is broken.",
        "Check for duplicate, incomplete, or corrupted response records and keep a clean export.",
        "Continue recruitment and monitor response count against the target agreed with the professor.",
        "The pilot has sufficient clean, usable responses or a clear plan to reach the target.",
    )
    add_week(
        doc,
        12,
        "Close and prepare the data",
        "Freeze feature work and only correct critical bugs.",
        "Export the final response dataset and lock the code/data version used for the pilot.",
        "Remove test records, document exclusions, and prepare the analysis-ready dataset.",
        "The team has a clean, frozen dataset and knows exactly what will be analysed.",
    )

    doc.add_page_break()
    add_heading(doc, "Month 4: Analysis, report, and final demo", level=1)
    add_week(
        doc,
        13,
        "Analyze the game results",
        "Capture polished screenshots and help convert key findings into clear visuals.",
        "Generate tables for accuracy, guesses, response time, and most-misclassified foods.",
        "Lead analysis by NOVA level and identify the foods players most often misunderstand.",
        "A results draft exists with accurate tables and a small set of meaningful findings.",
    )
    add_week(
        doc,
        14,
        "Write the report and presentation",
        "Prepare demo flow, screenshots, and visual consistency for the presentation.",
        "Document the database structure, deployment steps, and reproducible data workflow.",
        "Write methods, results, limitations, data attribution, and conclusion for the report.",
        "A complete first draft of the report and presentation exists.",
    )
    add_week(
        doc,
        15,
        "Final quality review",
        "Perform a full gameplay test on mobile and desktop and rehearse the demo.",
        "Check deployment, data export, and recovery of the final version.",
        "Check the report for terminology, source attribution, and consistency with the actual data.",
        "The app, report, and slides are internally approved by all three members.",
    )
    add_week(
        doc,
        16,
        "Submission buffer and presentation",
        "Handle final visual fixes and present the live game confidently.",
        "Be available for database and deployment questions during the demo.",
        "Be available for methodology, data, and analysis questions during the demo.",
        "The final deliverables are submitted, deployed, and ready to demonstrate.",
    )

    add_heading(doc, "This week's checklist: Week 2", level=1)
    add_callout(
        doc,
        "Keep this week focused.",
        "The goal is not a finished product. The goal is an approved initial card pool, a clickable static game screen, and a database structure ready to receive real cards.",
    )
    add_text_paragraph(doc, "Person A: build a static 10-question screen with four answer buttons and a score counter.", before=0, after=4, size=10.5)
    add_text_paragraph(doc, "Person B: create the foods, sessions, and responses tables plus a simple food-card import path.", before=0, after=4, size=10.5)
    add_text_paragraph(doc, "Person C: lead review of the 200 candidate cards. Each team member reviews roughly 65-70 cards and records approve/reject notes.", before=0, after=4, size=10.5)
    add_text_paragraph(doc, "All three: meet at the end of the week, combine approved cards, and test one 10-question round together.", before=0, after=4, size=10.5)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
