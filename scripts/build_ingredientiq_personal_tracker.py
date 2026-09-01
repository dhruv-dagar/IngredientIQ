from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "IngredientIQ_Personal_16_Week_Tracker.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"

TABLE_WIDTH = 14400
TABLE_INDENT = 90
WEEK_WIDTHS = [720, 1800, 3000, 3000, 3000, 1080, 1800]


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, before=0, after=0, line=1.08, keep_with_next=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=90, bottom=100, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths=WEEK_WIDTHS, indent=TABLE_INDENT):
    total = sum(widths)
    if total != TABLE_WIDTH:
        raise ValueError(f"Table widths must total {TABLE_WIDTH}; got {total}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for name, width, kind in (("tblW", total, "dxa"), ("tblInd", indent, "dxa")):
        element = tbl_pr.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tbl_pr.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), kind)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_text(cell, text, size=8.6, color=BLACK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_format(p, before=0, after=0, line=1.07)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instruction)
    run._r.append(fld_char2)
    return run


def configure_section(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.50)
    section.right_margin = Inches(0.50)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(header, before=0, after=0, line=1.0)
    r = header.add_run("INGREDIENTIQ | PERSONAL WEEKLY TRACKER")
    set_run_font(r, size=8.2, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(footer, before=0, after=0, line=1.0)
    r = footer.add_run("Private working tracker | Page ")
    set_run_font(r, size=8.2, color=MUTED)
    page = add_page_field(footer)
    set_run_font(page, size=8.2, color=MUTED)


def set_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=4, after=2, line=1.0, keep_with_next=True)
    r = p.add_run("INGREDIENTIQ PERSONAL TRACKER")
    set_run_font(r, size=22, color=NAVY, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=10, line=1.0, keep_with_next=True)
    r = p.add_run("16-week private working file for tracking the group project")
    set_run_font(r, size=11, color=MUTED, italic=True)

    info = doc.add_table(rows=1, cols=2)
    widths = [2880, 11520]
    set_table_geometry(info, widths=widths, indent=90)
    mark_header_row(info.rows[0])
    set_cell_text(info.cell(0, 0), "How to use", size=9.5, color=NAVY, bold=True)
    set_cell_text(
        info.cell(0, 1),
        "Type X inside [ ] as tasks are completed, or replace it with a tick mark. Use the Notes / blocker column during your weekly team check-in.",
        size=9.5,
    )
    for cell in info.rows[0].cells:
        shade_cell(cell, PALE_BLUE)
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, before=0, after=4, line=1.0)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=5, line=1.08)
    r = p.add_run("Personal note: ")
    set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    r = p.add_run("This is a working tracker, not a professor-facing document. It is not password-protected, so store it privately if needed.")
    set_run_font(r, size=9.5, color=BLACK)


def add_month_title(doc, title, target):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=5, line=1.0, keep_with_next=True)
    r = p.add_run(title)
    set_run_font(r, size=14, color=BLUE, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=7, line=1.0, keep_with_next=True)
    r = p.add_run("Month target: ")
    set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    r = p.add_run(target)
    set_run_font(r, size=9.5, color=BLACK)


def add_month_table(doc, weeks):
    table = doc.add_table(rows=1, cols=7)
    set_table_geometry(table)
    mark_header_row(table.rows[0])
    headers = [
        "Week",
        "Main aim",
        "A: Frontend",
        "B: Backend / DB",
        "C: Data / Research",
        "Track",
        "Notes / blocker",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, text, size=8.5, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for item in weeks:
        cells = table.add_row().cells
        values = [
            f"W{item['week']}",
            item["aim"],
            item["a"],
            item["b"],
            item["c"],
            "A [ ]\nB [ ]\nC [ ]\nWeek [ ]",
            "",
        ]
        for index, (cell, text) in enumerate(zip(cells, values)):
            if index == 0:
                shade_cell(cell, PALE_BLUE)
                set_cell_text(cell, text, size=9, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif index == 5:
                shade_cell(cell, PALE_BLUE)
                set_cell_text(cell, text, size=8.4, color=DARK_BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif index == 6:
                set_cell_text(cell, "\n\n\n", size=8.6)
            else:
                set_cell_text(cell, text, size=8.6)
    return table


MONTHS = [
    (
        "Month 1: Foundation and core game",
        "A local 10-question game works with approved, balanced food data.",
        [
            {
                "week": 1,
                "aim": "Scope and data setup",
                "a": "Map one round: card -> answer -> reveal -> score. Agree on minimum screens.",
                "b": "Choose stack; create repository; define foods, sessions, and responses fields.",
                "c": "Confirm NOVA terminology, inspect dataset, and write card-approval rules.",
            },
            {
                "week": 2,
                "aim": "Approved cards + static prototype",
                "a": "Build a static 10-question screen with four answers and score counter.",
                "b": "Create database schema and a repeatable food-card import path.",
                "c": "Lead review of 200 candidates. Team approves/rejects cards with notes.",
            },
            {
                "week": 3,
                "aim": "Connect real food data",
                "a": "Replace temporary cards with one random real food card at a time.",
                "b": "Seed approved foods and return questions without exposing the answer.",
                "c": "Check names, ingredients, duplicates, and NOVA-level balance.",
            },
            {
                "week": 4,
                "aim": "First complete game loop",
                "a": "Add feedback, reveal, score summary, and restart behaviour.",
                "b": "Save guess, actual level, response time, and anonymous session ID.",
                "c": "Write factual reveal text and check every seeded card.",
            },
        ],
    ),
    (
        "Month 2: MVP completion and internal testing",
        "The game is mobile-ready, shareable, and stable enough for a professor demo.",
        [
            {
                "week": 5,
                "aim": "Mobile reliability",
                "a": "Improve responsive layout, loading state, and first-time instructions.",
                "b": "Prevent repeats in a round and validate answer-logging rules.",
                "c": "Replace weak/unfamiliar cards; keep 100+ balanced approved cards.",
            },
            {
                "week": 6,
                "aim": "Minimal data collection",
                "a": "Improve start, results, and replay screens; keep design simple.",
                "b": "Add anonymous sessions and response export. Skip login unless required.",
                "c": "Define analysis metrics: accuracy, wrong guesses, score, and time.",
            },
            {
                "week": 7,
                "aim": "Staging and usability test",
                "a": "Polish labels, button states, accessibility, and mobile spacing.",
                "b": "Deploy a staging link and test data flow off the local computer.",
                "c": "Run 5-10 person usability test; record clear feedback.",
            },
            {
                "week": 8,
                "aim": "MVP checkpoint",
                "a": "Fix important usability issues; do not add game modes.",
                "b": "Fix deployment, data, and logging issues found in testing.",
                "c": "Update cards/wording and prepare professor checkpoint notes.",
            },
        ],
    ),
    (
        "Month 3: Pilot preparation and data collection",
        "Collect a clean, stable pilot dataset without changing core game behaviour mid-study.",
        [
            {
                "week": 9,
                "aim": "Prepare the pilot",
                "a": "Make only small clarity fixes and a short instruction screen.",
                "b": "Add simple card-update workflow only if currently necessary.",
                "c": "Prepare participant instructions, recruitment message, and consent text if needed.",
            },
            {
                "week": 10,
                "aim": "Launch pilot",
                "a": "Support players and fix only game-breaking issues.",
                "b": "Monitor errors; confirm response data stores and exports correctly.",
                "c": "Start recruitment; track participation and card complaints.",
            },
            {
                "week": 11,
                "aim": "Continue data collection",
                "a": "Keep UI and game logic stable while the pilot runs.",
                "b": "Check for duplicate, incomplete, or corrupted response records.",
                "c": "Continue recruitment and compare response count with the target.",
            },
            {
                "week": 12,
                "aim": "Close and prepare data",
                "a": "Freeze feature work; fix critical bugs only.",
                "b": "Export final responses and lock code/data used for the pilot.",
                "c": "Remove test records, document exclusions, and prepare analysis file.",
            },
        ],
    ),
    (
        "Month 4: Analysis, report, and final demo",
        "Deliver a polished game, analysis, report, and presentation with time left for fixes.",
        [
            {
                "week": 13,
                "aim": "Analyze results",
                "a": "Capture polished screenshots and help present key findings clearly.",
                "b": "Generate tables for accuracy, guesses, time, and misclassified foods.",
                "c": "Analyze results by NOVA level and identify misunderstood foods.",
            },
            {
                "week": 14,
                "aim": "Report and slides",
                "a": "Prepare demo flow, screenshots, and visual consistency.",
                "b": "Document database structure, deployment, and data workflow.",
                "c": "Write methods, results, limitations, attribution, and conclusion.",
            },
            {
                "week": 15,
                "aim": "Final quality review",
                "a": "Test full game on mobile/desktop and rehearse demo.",
                "b": "Check deployment, exports, and final-version recovery.",
                "c": "Check report terminology, source attribution, and data consistency.",
            },
            {
                "week": 16,
                "aim": "Submission buffer + presentation",
                "a": "Handle final visual fixes and present the live game.",
                "b": "Answer database and deployment questions during the demo.",
                "c": "Answer methodology, data, and analysis questions during the demo.",
            },
        ],
    ),
]


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc.sections[0])
    set_document_styles(doc)
    add_title_block(doc)

    for index, (title, target, weeks) in enumerate(MONTHS):
        if index > 0:
            doc.add_page_break()
        add_month_title(doc, title, target)
        add_month_table(doc, weeks)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
